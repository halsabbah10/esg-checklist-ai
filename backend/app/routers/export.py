import logging
from datetime import datetime
from io import BytesIO, StringIO
from typing import Any

import pandas as pd
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from fastapi import APIRouter, Depends, HTTPException, Query, Request
from fastapi.responses import StreamingResponse
from sqlmodel import Session, select

# Import for Word document generation
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.auth import require_role
from app.database import get_session
from app.models import AIResult, Checklist, FileUpload, SubmissionAnswer, User, UserActivity, SystemMetrics
from app.rate_limiting import admin_rate_limit, export_rate_limit

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/export", tags=["export"])


@router.get("/checklists")
@export_rate_limit
def export_all_checklists(
    request: Request,
    format: str = Query("csv", pattern="^(csv|excel|json|pdf|docx)$"),
    include_inactive: bool = Query(False, description="Include inactive checklists"),
    db: Session = Depends(get_session),
    current_user=Depends(require_role(["admin", "reviewer"])),
):
    """Export all checklists in various formats"""
    try:
        # Build query
        query: Any = select(Checklist)
        if not include_inactive:
            query = query.where(Checklist.is_active is True)

        checklists = db.exec(query).all()

        # Prepare data
        data = []
        for checklist in checklists:
            data.append(
                {
                    "id": checklist.id,
                    "title": checklist.title,
                    "description": checklist.description,
                    "created_by": checklist.created_by,
                    "created_at": checklist.created_at,
                    "updated_at": getattr(checklist, "updated_at", None),  # Safe access
                    "is_active": checklist.is_active,
                    "version": checklist.version,
                }
            )

        df = pd.DataFrame(data)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        if format == "csv":
            csv_buf = StringIO()
            df.to_csv(csv_buf, index=False)
            csv_buf.seek(0)
            return StreamingResponse(
                iter([csv_buf.getvalue()]),
                media_type="text/csv",
                headers={"Content-Disposition": f"attachment; filename=checklists_{timestamp}.csv"},
            )
        if format == "excel":
            excel_buf = BytesIO()
            df.to_excel(excel_buf, index=False, engine="openpyxl")
            excel_buf.seek(0)
            return StreamingResponse(
                iter([excel_buf.getvalue()]),
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={
                    "Content-Disposition": f"attachment; filename=checklists_{timestamp}.xlsx"
                },
            )
        if format == "json":
            json_buf = StringIO()
            df.to_json(json_buf, orient="records", indent=2)
            json_buf.seek(0)
            return StreamingResponse(
                iter([json_buf.getvalue()]),
                media_type="application/json",
                headers={
                    "Content-Disposition": f"attachment; filename=checklists_{timestamp}.json"
                },
            )
        if format == "pdf":
            pdf_buf = BytesIO()
            doc = SimpleDocTemplate(pdf_buf, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.HexColor('#1976d2')
            )
            story.append(Paragraph("ESG Checklists Export Report", title_style))
            story.append(Spacer(1, 12))
            
            # Summary
            summary_style = styles['Normal']
            story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", summary_style))
            story.append(Paragraph(f"Total Checklists: {len(data)}", summary_style))
            story.append(Spacer(1, 20))
            
            # Data table
            if data:
                table_data = [['ID', 'Title', 'Description', 'Status', 'Created']]
                for row in data:
                    table_data.append([
                        str(row['id']),
                        row['title'][:30] + '...' if len(row['title']) > 30 else row['title'],
                        row['description'][:40] + '...' if row['description'] and len(row['description']) > 40 else (row['description'] or ''),
                        'Active' if row['is_active'] else 'Inactive',
                        row['created_at'].strftime('%Y-%m-%d') if row['created_at'] else ''
                    ])
                
                table = Table(table_data, colWidths=[0.8*inch, 2*inch, 2.5*inch, 1*inch, 1*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1976d2')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
            
            doc.build(story)
            pdf_buf.seek(0)
            return StreamingResponse(
                iter([pdf_buf.getvalue()]),
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f"attachment; filename=checklists_{timestamp}.pdf"
                },
            )
        if format == "docx":
            if not DOCX_AVAILABLE:
                raise HTTPException(status_code=500, detail="Word document generation not available. Missing python-docx dependency.")
            
            docx_buf = BytesIO()
            doc = Document()
            
            # Title
            title = doc.add_heading('ESG Checklists Export Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Summary
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Total Checklists: {len(data)}")
            doc.add_paragraph()
            
            # Data table
            if data:
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Light Grid Accent 1'
                
                # Header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'ID'
                hdr_cells[1].text = 'Title'
                hdr_cells[2].text = 'Description'
                hdr_cells[3].text = 'Status'
                hdr_cells[4].text = 'Created'
                
                # Data rows
                for row_data in data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row_data['id'])
                    row_cells[1].text = row_data['title'][:40] + '...' if len(row_data['title']) > 40 else row_data['title']
                    row_cells[2].text = (row_data['description'] or '')[:50] + '...' if row_data['description'] and len(row_data['description']) > 50 else (row_data['description'] or '')
                    row_cells[3].text = 'Active' if row_data['is_active'] else 'Inactive'
                    row_cells[4].text = row_data['created_at'].strftime('%Y-%m-%d') if row_data['created_at'] else ''
            
            doc.save(docx_buf)
            docx_buf.seek(0)
            return StreamingResponse(
                iter([docx_buf.getvalue()]),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f"attachment; filename=checklists_{timestamp}.docx"
                },
            )

    except Exception as e:
        logger.exception(f"Export error: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {e!s}")


@router.get("/ai-results")
@export_rate_limit
def export_ai_results(
    request: Request,
    format: str = Query("csv", pattern="^(csv|excel|json|pdf|docx)$"),
    checklist_id: int = Query(None, description="Filter by checklist ID"),
    min_score: float = Query(None, ge=0.0, le=1.0, description="Minimum AI score"),
    days: int = Query(30, ge=1, le=365, description="Days of data to include"),
    db: Session = Depends(get_session),
    current_user=Depends(require_role(["admin", "reviewer"])),
):
    """Export AI scoring results with filtering options"""
    try:
        # Build query with joins
        query: Any = (
            select(AIResult, FileUpload, Checklist, User)
            .join(FileUpload, AIResult.file_upload_id == FileUpload.id)
            .join(Checklist, AIResult.checklist_id == Checklist.id)
            .join(User, AIResult.user_id == User.id)
        )

        # Apply filters
        if checklist_id:
            query = query.where(AIResult.checklist_id == checklist_id)
        if min_score is not None:
            query = query.where(AIResult.score >= min_score)

        # Date filter
        cutoff_date = datetime.now() - pd.Timedelta(days=days)
        query = query.where(AIResult.created_at >= cutoff_date)

        results = db.exec(query).all()

        # Prepare data
        data = []
        for ai_result, file_upload, checklist, user in results:
            data.append(
                {
                    "ai_result_id": ai_result.id,
                    "checklist_id": checklist.id,
                    "checklist_title": checklist.title,
                    "file_id": file_upload.id,
                    "filename": file_upload.filename,
                    "user_id": user.id,
                    "username": user.username,
                    "user_email": user.email,
                    "ai_score": ai_result.score,
                    "feedback": (
                        ai_result.feedback[:500] + "..."
                        if len(ai_result.feedback) > 500
                        else ai_result.feedback
                    ),
                    "processing_time_ms": ai_result.processing_time_ms,
                    "ai_model_version": ai_result.ai_model_version,
                    "created_at": ai_result.created_at,
                    "uploaded_at": file_upload.uploaded_at,
                }
            )

        df = pd.DataFrame(data)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        if format == "csv":
            csv_buf = StringIO()
            df.to_csv(csv_buf, index=False)
            csv_buf.seek(0)
            return StreamingResponse(
                iter([csv_buf.getvalue()]),
                media_type="text/csv",
                headers={"Content-Disposition": f"attachment; filename=ai_results_{timestamp}.csv"},
            )
        if format == "excel":
            excel_buf = BytesIO()
            df.to_excel(excel_buf, index=False, engine="openpyxl")
            excel_buf.seek(0)
            return StreamingResponse(
                iter([excel_buf.getvalue()]),
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={
                    "Content-Disposition": f"attachment; filename=ai_results_{timestamp}.xlsx"
                },
            )
        if format == "json":
            json_buf = StringIO()
            df.to_json(json_buf, orient="records", indent=2)
            json_buf.seek(0)
            return StreamingResponse(
                iter([json_buf.getvalue()]),
                media_type="application/json",
                headers={
                    "Content-Disposition": f"attachment; filename=ai_results_{timestamp}.json"
                },
            )
        if format == "pdf":
            pdf_buf = BytesIO()
            doc = SimpleDocTemplate(pdf_buf, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.HexColor('#1976d2')
            )
            story.append(Paragraph("AI Analysis Results Export Report", title_style))
            story.append(Spacer(1, 12))
            
            # Summary
            summary_style = styles['Normal']
            story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", summary_style))
            story.append(Paragraph(f"Total AI Results: {len(data)}", summary_style))
            if checklist_id:
                story.append(Paragraph(f"Filtered by Checklist ID: {checklist_id}", summary_style))
            if min_score is not None:
                story.append(Paragraph(f"Minimum Score Filter: {min_score}", summary_style))
            story.append(Spacer(1, 20))
            
            # Data table
            if data:
                table_data = [['File', 'User', 'Score', 'Checklist', 'Date']]
                for row in data:
                    table_data.append([
                        row['filename'][:25] + '...' if len(row['filename']) > 25 else row['filename'],
                        row['username'][:15] + '...' if len(row['username']) > 15 else row['username'],
                        f"{row['ai_score']:.1%}" if row['ai_score'] is not None else 'N/A',
                        row['checklist_title'][:20] + '...' if len(row['checklist_title']) > 20 else row['checklist_title'],
                        row['created_at'].strftime('%Y-%m-%d') if row['created_at'] else ''
                    ])
                
                table = Table(table_data, colWidths=[2*inch, 1.5*inch, 1*inch, 1.5*inch, 1*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1976d2')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
            
            doc.build(story)
            pdf_buf.seek(0)
            return StreamingResponse(
                iter([pdf_buf.getvalue()]),
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f"attachment; filename=ai_results_{timestamp}.pdf"
                },
            )
        if format == "docx":
            if not DOCX_AVAILABLE:
                raise HTTPException(status_code=500, detail="Word document generation not available. Missing python-docx dependency.")
            
            docx_buf = BytesIO()
            doc = Document()
            
            # Title
            title = doc.add_heading('AI Analysis Results Export Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Summary
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Total AI Results: {len(data)}")
            if checklist_id:
                doc.add_paragraph(f"Filtered by Checklist ID: {checklist_id}")
            if min_score is not None:
                doc.add_paragraph(f"Minimum Score Filter: {min_score}")
            doc.add_paragraph()
            
            # Data table
            if data:
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Light Grid Accent 1'
                
                # Header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'File'
                hdr_cells[1].text = 'User'
                hdr_cells[2].text = 'Score'
                hdr_cells[3].text = 'Checklist'
                hdr_cells[4].text = 'Date'
                
                # Data rows
                for row_data in data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = row_data['filename'][:30] + '...' if len(row_data['filename']) > 30 else row_data['filename']
                    row_cells[1].text = row_data['username'][:20] + '...' if len(row_data['username']) > 20 else row_data['username']
                    row_cells[2].text = f"{row_data['ai_score']:.1%}" if row_data['ai_score'] is not None else 'N/A'
                    row_cells[3].text = row_data['checklist_title'][:25] + '...' if len(row_data['checklist_title']) > 25 else row_data['checklist_title']
                    row_cells[4].text = row_data['created_at'].strftime('%Y-%m-%d') if row_data['created_at'] else ''
            
            doc.save(docx_buf)
            docx_buf.seek(0)
            return StreamingResponse(
                iter([docx_buf.getvalue()]),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f"attachment; filename=ai_results_{timestamp}.docx"
                },
            )

    except Exception as e:
        logger.exception(f"AI results export error: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {e!s}")


@router.get("/users")
@admin_rate_limit
def export_users(
    request: Request,
    format: str = Query("csv", pattern="^(csv|excel|json|pdf|docx)$"),
    role: str = Query(None, description="Filter by user role"),
    include_stats: bool = Query(True, description="Include user activity statistics"),
    db: Session = Depends(get_session),
    current_user=Depends(require_role(["admin", "reviewer"])),
):
    """Export user data with optional statistics"""
    try:
        # Get users
        query: Any = select(User)
        if role:
            query = query.where(User.role == role)
        users = db.exec(query).all()

        # Prepare data
        data = []
        for user in users:
            user_data = {
                "id": user.id,
                "username": user.username,
                "email": user.email,
                "role": user.role,
                "created_at": user.created_at,
                "last_login": user.last_login,
                "is_active": user.is_active,
            }

            # Add statistics if requested
            if include_stats:
                file_count = db.exec(select(FileUpload).where(FileUpload.user_id == user.id)).all()
                ai_results = db.exec(select(AIResult).where(AIResult.user_id == user.id)).all()

                user_data.update(
                    {
                        "total_uploads": len(file_count),
                        "total_ai_analyses": len(ai_results),
                        "avg_ai_score": (
                            sum(r.score for r in ai_results) / len(ai_results)
                            if ai_results
                            else 0.0
                        ),
                        "last_activity": (
                            max([f.uploaded_at for f in file_count] + [user.created_at])
                            if file_count
                            else user.created_at
                        ),
                    }
                )

            data.append(user_data)

        df = pd.DataFrame(data)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        if format == "csv":
            csv_buf = StringIO()
            df.to_csv(csv_buf, index=False)
            csv_buf.seek(0)
            return StreamingResponse(
                iter([csv_buf.getvalue()]),
                media_type="text/csv",
                headers={"Content-Disposition": f"attachment; filename=users_{timestamp}.csv"},
            )
        if format == "excel":
            excel_buf = BytesIO()
            df.to_excel(excel_buf, index=False, engine="openpyxl")
            excel_buf.seek(0)
            return StreamingResponse(
                iter([excel_buf.getvalue()]),
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": f"attachment; filename=users_{timestamp}.xlsx"},
            )
        if format == "json":
            json_buf = StringIO()
            df.to_json(json_buf, orient="records", indent=2)
            json_buf.seek(0)
            return StreamingResponse(
                iter([json_buf.getvalue()]),
                media_type="application/json",
                headers={"Content-Disposition": f"attachment; filename=users_{timestamp}.json"},
            )
        if format == "pdf":
            pdf_buf = BytesIO()
            doc = SimpleDocTemplate(pdf_buf, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.HexColor('#1976d2')
            )
            story.append(Paragraph("Users Directory Export Report", title_style))
            story.append(Spacer(1, 12))
            
            # Summary
            summary_style = styles['Normal']
            story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", summary_style))
            story.append(Paragraph(f"Total Users: {len(data)}", summary_style))
            if role:
                story.append(Paragraph(f"Filtered by Role: {role}", summary_style))
            story.append(Spacer(1, 20))
            
            # Data table
            if data:
                if include_stats:
                    table_data = [['Username', 'Email', 'Role', 'Uploads', 'Avg Score', 'Status']]
                    for row in data:
                        table_data.append([
                            row['username'][:20] + '...' if len(row['username']) > 20 else row['username'],
                            row['email'][:25] + '...' if len(row['email']) > 25 else row['email'],
                            row['role'].title(),
                            str(row.get('total_uploads', 0)),
                            f"{row.get('avg_ai_score', 0):.1%}" if row.get('avg_ai_score') else 'N/A',
                            'Active' if row['is_active'] else 'Inactive'
                        ])
                else:
                    table_data = [['Username', 'Email', 'Role', 'Created', 'Status']]
                    for row in data:
                        table_data.append([
                            row['username'][:20] + '...' if len(row['username']) > 20 else row['username'],
                            row['email'][:30] + '...' if len(row['email']) > 30 else row['email'],
                            row['role'].title(),
                            row['created_at'].strftime('%Y-%m-%d') if row['created_at'] else '',
                            'Active' if row['is_active'] else 'Inactive'
                        ])
                
                table = Table(table_data, colWidths=[1.5*inch, 2*inch, 1*inch, 1*inch, 1*inch, 0.8*inch] if include_stats else [1.8*inch, 2.5*inch, 1*inch, 1.2*inch, 1*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1976d2')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
            
            doc.build(story)
            pdf_buf.seek(0)
            return StreamingResponse(
                iter([pdf_buf.getvalue()]),
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f"attachment; filename=users_{timestamp}.pdf"
                },
            )

    except Exception as e:
        logger.exception(f"Users export error: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {e!s}")


@router.get("/submissions")
@export_rate_limit
def export_submissions(
    request: Request,
    format: str = Query("csv", pattern="^(csv|excel|json|pdf|docx)$"),
    checklist_id: int = Query(None, description="Filter by checklist ID"),
    user_id: int = Query(None, description="Filter by user ID"),
    days: int = Query(30, ge=1, le=365, description="Days of data to include"),
    db: Session = Depends(get_session),
    current_user=Depends(require_role(["admin", "reviewer"])),
):
    """Export submission answers with filtering options"""
    try:
        # Build query with joins
        query: Any = (
            select(SubmissionAnswer, Checklist, User)
            .join(
                Checklist, SubmissionAnswer.checklist_id == Checklist.id
            )
            .join(User, SubmissionAnswer.user_id == User.id)
        )

        # Apply filters
        if checklist_id:
            query = query.where(SubmissionAnswer.checklist_id == checklist_id)
        if user_id:
            query = query.where(SubmissionAnswer.user_id == user_id)

        # Date filter
        cutoff_date = datetime.now() - pd.Timedelta(days=days)
        query = query.where(SubmissionAnswer.submitted_at >= cutoff_date)

        results = db.exec(query).all()

        # Prepare data
        data = []
        for submission, checklist, user in results:
            data.append(
                {
                    "submission_id": submission.id,
                    "checklist_id": checklist.id,
                    "checklist_title": checklist.title,
                    "question_id": submission.question_id,
                    "user_id": user.id,
                    "username": user.username,
                    "user_email": user.email,
                    "user_role": user.role,
                    "answer_text": (
                        submission.answer_text[:1000] + "..."
                        if len(submission.answer_text) > 1000
                        else submission.answer_text
                    ),
                    "submitted_at": submission.submitted_at,
                }
            )

        df = pd.DataFrame(data)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        if format == "csv":
            csv_buf = StringIO()
            df.to_csv(csv_buf, index=False)
            csv_buf.seek(0)
            return StreamingResponse(
                iter([csv_buf.getvalue()]),
                media_type="text/csv",
                headers={
                    "Content-Disposition": f"attachment; filename=submissions_{timestamp}.csv"
                },
            )
        if format == "excel":
            excel_buf = BytesIO()
            df.to_excel(excel_buf, index=False, engine="openpyxl")
            excel_buf.seek(0)
            return StreamingResponse(
                iter([excel_buf.getvalue()]),
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={
                    "Content-Disposition": f"attachment; filename=submissions_{timestamp}.xlsx"
                },
            )
        if format == "json":
            json_buf = StringIO()
            df.to_json(json_buf, orient="records", indent=2)
            json_buf.seek(0)
            return StreamingResponse(
                iter([json_buf.getvalue()]),
                media_type="application/json",
                headers={
                    "Content-Disposition": f"attachment; filename=submissions_{timestamp}.json"
                },
            )
        if format == "pdf":
            pdf_buf = BytesIO()
            doc = SimpleDocTemplate(pdf_buf, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.HexColor('#1976d2')
            )
            story.append(Paragraph("User Submissions Export Report", title_style))
            story.append(Spacer(1, 12))
            
            # Summary
            summary_style = styles['Normal']
            story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", summary_style))
            story.append(Paragraph(f"Total Submissions: {len(data)}", summary_style))
            if checklist_id:
                story.append(Paragraph(f"Filtered by Checklist ID: {checklist_id}", summary_style))
            if user_id:
                story.append(Paragraph(f"Filtered by User ID: {user_id}", summary_style))
            story.append(Spacer(1, 20))
            
            # Data table
            if data:
                table_data = [['User', 'Checklist', 'Question ID', 'Answer Preview', 'Date']]
                for row in data:
                    answer_preview = row['answer_text'][:40] + '...' if len(row['answer_text']) > 40 else row['answer_text']
                    table_data.append([
                        row['username'][:15] + '...' if len(row['username']) > 15 else row['username'],
                        row['checklist_title'][:20] + '...' if len(row['checklist_title']) > 20 else row['checklist_title'],
                        str(row['question_id']),
                        answer_preview,
                        row['submitted_at'].strftime('%Y-%m-%d') if row['submitted_at'] else ''
                    ])
                
                table = Table(table_data, colWidths=[1.2*inch, 1.8*inch, 0.8*inch, 2.5*inch, 1*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1976d2')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
            
            doc.build(story)
            pdf_buf.seek(0)
            return StreamingResponse(
                iter([pdf_buf.getvalue()]),
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f"attachment; filename=submissions_{timestamp}.pdf"
                },
            )

    except Exception as e:
        logger.exception(f"Submissions export error: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {e!s}")


@router.get("/analytics")
@export_rate_limit
def export_analytics_data(
    request: Request,
    format: str = Query("csv", pattern="^(csv|excel|json|pdf|docx)$"),
    days: int = Query(30, ge=1, le=365, description="Days of data to include"),
    include_scores: bool = Query(True, description="Include AI scores in export"),
    db: Session = Depends(get_session),
    current_user=Depends(require_role(["admin", "reviewer"])),
):
    """Export comprehensive analytics data"""
    try:
        from datetime import datetime, timedelta
        
        cutoff_date = datetime.now() - timedelta(days=days)
        
        # Get comprehensive analytics data
        data = []
        
        # Get users with their activity
        users = db.exec(select(User)).all()
        for user in users:
            # Get user's uploads
            user_uploads = db.exec(
                select(FileUpload).where(
                    FileUpload.user_id == user.id,
                    FileUpload.uploaded_at >= cutoff_date
                )
            ).all()
            
            # Get user's AI results
            user_ai_results = db.exec(
                select(AIResult).where(
                    AIResult.user_id == user.id,
                    AIResult.created_at >= cutoff_date
                )
            ).all()
            
            # Get user activities
            user_activities = db.exec(
                select(UserActivity).where(
                    UserActivity.user_id == user.id,
                    UserActivity.timestamp >= cutoff_date
                )
            ).all()
            
            # Calculate metrics
            avg_score = (
                sum(r.score for r in user_ai_results) / len(user_ai_results)
                if user_ai_results and include_scores
                else 0
            )
            
            avg_processing_time = (
                sum(r.processing_time_ms or 0 for r in user_ai_results) / len(user_ai_results)
                if user_ai_results
                else 0
            )
            
            data.append({
                "user_id": user.id,
                "username": user.username,
                "email": user.email,
                "role": user.role,
                "created_at": user.created_at,
                "last_login": user.last_login,
                "is_active": user.is_active,
                "total_uploads": len(user_uploads),
                "total_ai_analyses": len(user_ai_results),
                "avg_ai_score": round(avg_score, 3) if include_scores else None,
                "avg_processing_time_ms": round(avg_processing_time, 2),
                "total_activities": len(user_activities),
                "most_recent_activity": max([a.timestamp for a in user_activities]) if user_activities else None,
                "activity_types": list(set(a.action_type for a in user_activities)),
            })
        
        df = pd.DataFrame(data)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if format == "csv":
            csv_buf = StringIO()
            df.to_csv(csv_buf, index=False)
            csv_buf.seek(0)
            return StreamingResponse(
                iter([csv_buf.getvalue()]),
                media_type="text/csv",
                headers={"Content-Disposition": f"attachment; filename=analytics_{timestamp}.csv"},
            )
        elif format == "excel":
            excel_buf = BytesIO()
            df.to_excel(excel_buf, index=False, engine="openpyxl")
            excel_buf.seek(0)
            return StreamingResponse(
                iter([excel_buf.getvalue()]),
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": f"attachment; filename=analytics_{timestamp}.xlsx"},
            )
        elif format == "json":
            json_buf = StringIO()
            df.to_json(json_buf, orient="records", indent=2)
            json_buf.seek(0)
            return StreamingResponse(
                iter([json_buf.getvalue()]),
                media_type="application/json",
                headers={"Content-Disposition": f"attachment; filename=analytics_{timestamp}.json"},
            )
        elif format == "pdf":
            pdf_buf = BytesIO()
            doc = SimpleDocTemplate(pdf_buf, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.HexColor('#1976d2')
            )
            story.append(Paragraph("Analytics Export Report", title_style))
            story.append(Spacer(1, 12))
            
            # Summary
            summary_style = styles['Normal']
            story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", summary_style))
            story.append(Paragraph(f"Total Users Analyzed: {len(data)}", summary_style))
            story.append(Paragraph(f"Data Period: Last {days} days", summary_style))
            story.append(Spacer(1, 20))
            
            # Data table
            if data:
                if include_scores:
                    table_data = [['User', 'Role', 'Uploads', 'AI Analyses', 'Avg Score', 'Activities']]
                    for row in data:
                        table_data.append([
                            row['username'][:15] + '...' if len(row['username']) > 15 else row['username'],
                            row['role'].title(),
                            str(row['total_uploads']),
                            str(row['total_ai_analyses']),
                            f"{row['avg_ai_score']:.1%}" if row['avg_ai_score'] else 'N/A',
                            str(row['total_activities'])
                        ])
                else:
                    table_data = [['User', 'Role', 'Uploads', 'AI Analyses', 'Activities', 'Status']]
                    for row in data:
                        table_data.append([
                            row['username'][:15] + '...' if len(row['username']) > 15 else row['username'],
                            row['role'].title(),
                            str(row['total_uploads']),
                            str(row['total_ai_analyses']),
                            str(row['total_activities']),
                            'Active' if row['is_active'] else 'Inactive'
                        ])
                
                table = Table(table_data, colWidths=[1.3*inch, 1*inch, 0.8*inch, 1*inch, 0.8*inch, 0.8*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1976d2')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
            
            doc.build(story)
            pdf_buf.seek(0)
            return StreamingResponse(
                iter([pdf_buf.getvalue()]),
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename=analytics_{timestamp}.pdf"},
            )
            
    except Exception as e:
        logger.exception(f"Analytics export error: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {e!s}")


@router.get("/system-metrics")
@export_rate_limit
def export_system_metrics(
    request: Request,
    format: str = Query("csv", pattern="^(csv|excel|json|pdf|docx)$"),
    days: int = Query(7, ge=1, le=365, description="Days of data to include"),
    category: str = Query(None, description="Filter by metric category"),
    db: Session = Depends(get_session),
    current_user=Depends(require_role(["admin", "reviewer"])),
):
    """Export system metrics and performance data"""
    try:
        from datetime import datetime, timedelta
        
        cutoff_date = datetime.now() - timedelta(days=days)
        
        # Build query
        query = select(SystemMetrics).where(SystemMetrics.recorded_at >= cutoff_date)
        if category:
            query = query.where(SystemMetrics.category == category)
        
        metrics = db.exec(query.order_by(SystemMetrics.recorded_at.desc())).all()
        
        # Prepare data
        data = []
        for metric in metrics:
            data.append({
                "id": metric.id,
                "metric_name": metric.metric_name,
                "metric_value": metric.metric_value,
                "metric_unit": metric.metric_unit,
                "category": metric.category,
                "recorded_at": metric.recorded_at,
                "additional_data": metric.additional_data,
            })
        
        df = pd.DataFrame(data)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if format == "csv":
            csv_buf = StringIO()
            df.to_csv(csv_buf, index=False)
            csv_buf.seek(0)
            return StreamingResponse(
                iter([csv_buf.getvalue()]),
                media_type="text/csv",
                headers={"Content-Disposition": f"attachment; filename=system_metrics_{timestamp}.csv"},
            )
        elif format == "excel":
            excel_buf = BytesIO()
            df.to_excel(excel_buf, index=False, engine="openpyxl")
            excel_buf.seek(0)
            return StreamingResponse(
                iter([excel_buf.getvalue()]),
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": f"attachment; filename=system_metrics_{timestamp}.xlsx"},
            )
        elif format == "json":
            json_buf = StringIO()
            df.to_json(json_buf, orient="records", indent=2)
            json_buf.seek(0)
            return StreamingResponse(
                iter([json_buf.getvalue()]),
                media_type="application/json",
                headers={"Content-Disposition": f"attachment; filename=system_metrics_{timestamp}.json"},
            )
        elif format == "pdf":
            pdf_buf = BytesIO()
            doc = SimpleDocTemplate(pdf_buf, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.HexColor('#1976d2')
            )
            story.append(Paragraph("System Metrics Export Report", title_style))
            story.append(Spacer(1, 12))
            
            # Summary
            summary_style = styles['Normal']
            story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", summary_style))
            story.append(Paragraph(f"Total Metrics: {len(data)}", summary_style))
            story.append(Paragraph(f"Data Period: Last {days} days", summary_style))
            if category:
                story.append(Paragraph(f"Category Filter: {category}", summary_style))
            story.append(Spacer(1, 20))
            
            # Data table
            if data:
                table_data = [['Metric Name', 'Value', 'Unit', 'Category', 'Recorded']]
                for row in data:
                    table_data.append([
                        row['metric_name'][:25] + '...' if len(row['metric_name']) > 25 else row['metric_name'],
                        str(row['metric_value'])[:15] + '...' if len(str(row['metric_value'])) > 15 else str(row['metric_value']),
                        row['metric_unit'] or '',
                        row['category'] or '',
                        row['recorded_at'].strftime('%Y-%m-%d %H:%M') if row['recorded_at'] else ''
                    ])
                
                table = Table(table_data, colWidths=[2*inch, 1.2*inch, 0.8*inch, 1.2*inch, 1.3*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1976d2')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
            
            doc.build(story)
            pdf_buf.seek(0)
            return StreamingResponse(
                iter([pdf_buf.getvalue()]),
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename=system_metrics_{timestamp}.pdf"},
            )
            
    except Exception as e:
        logger.exception(f"System metrics export error: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {e!s}")


@router.get("/user-activities")
@export_rate_limit
def export_user_activities(
    request: Request,
    format: str = Query("csv", pattern="^(csv|excel|json|pdf|docx)$"),
    days: int = Query(7, ge=1, le=30, description="Days of data to include"),
    user_id: int = Query(None, description="Filter by user ID"),
    action_type: str = Query(None, description="Filter by action type"),
    db: Session = Depends(get_session),
    current_user=Depends(require_role(["admin", "reviewer"])),
):
    """Export user activity logs"""
    try:
        from datetime import datetime, timedelta
        
        cutoff_date = datetime.now() - timedelta(days=days)
        
        # Build query with joins
        query = (
            select(UserActivity, User.username, User.email)
            .join(User, UserActivity.user_id == User.id)
            .where(UserActivity.timestamp >= cutoff_date)
        )
        
        if user_id:
            query = query.where(UserActivity.user_id == user_id)
        if action_type:
            query = query.where(UserActivity.action_type == action_type)
        
        activities = db.exec(query.order_by(UserActivity.timestamp.desc())).all()
        
        # Prepare data
        data = []
        for activity, username, email in activities:
            data.append({
                "id": activity.id,
                "user_id": activity.user_id,
                "username": username,
                "email": email,
                "session_id": activity.session_id,
                "action_type": activity.action_type,
                "resource_type": activity.resource_type,
                "resource_id": activity.resource_id,
                "duration_ms": activity.duration_ms,
                "ip_address": activity.ip_address,
                "user_agent": activity.user_agent,
                "timestamp": activity.timestamp,
                "action_details": activity.action_details,
            })
        
        df = pd.DataFrame(data)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if format == "csv":
            csv_buf = StringIO()
            df.to_csv(csv_buf, index=False)
            csv_buf.seek(0)
            return StreamingResponse(
                iter([csv_buf.getvalue()]),
                media_type="text/csv",
                headers={"Content-Disposition": f"attachment; filename=user_activities_{timestamp}.csv"},
            )
        elif format == "excel":
            excel_buf = BytesIO()
            df.to_excel(excel_buf, index=False, engine="openpyxl")
            excel_buf.seek(0)
            return StreamingResponse(
                iter([excel_buf.getvalue()]),
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": f"attachment; filename=user_activities_{timestamp}.xlsx"},
            )
        elif format == "json":
            json_buf = StringIO()
            df.to_json(json_buf, orient="records", indent=2)
            json_buf.seek(0)
            return StreamingResponse(
                iter([json_buf.getvalue()]),
                media_type="application/json",
                headers={"Content-Disposition": f"attachment; filename=user_activities_{timestamp}.json"},
            )
        elif format == "pdf":
            pdf_buf = BytesIO()
            doc = SimpleDocTemplate(pdf_buf, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.HexColor('#1976d2')
            )
            story.append(Paragraph("User Activities Export Report", title_style))
            story.append(Spacer(1, 12))
            
            # Summary
            summary_style = styles['Normal']
            story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", summary_style))
            story.append(Paragraph(f"Total Activities: {len(data)}", summary_style))
            story.append(Paragraph(f"Data Period: Last {days} days", summary_style))
            if user_id:
                story.append(Paragraph(f"User Filter: {user_id}", summary_style))
            if action_type:
                story.append(Paragraph(f"Action Type Filter: {action_type}", summary_style))
            story.append(Spacer(1, 20))
            
            # Data table
            if data:
                table_data = [['User', 'Action', 'Resource', 'Duration', 'Timestamp']]
                for row in data:
                    duration_ms = row['duration_ms'] or 0
                    duration_str = f"{duration_ms}ms" if duration_ms < 1000 else f"{duration_ms/1000:.1f}s"
                    table_data.append([
                        row['username'][:15] + '...' if len(row['username']) > 15 else row['username'],
                        row['action_type'][:15] + '...' if len(row['action_type']) > 15 else row['action_type'],
                        row['resource_type'][:15] + '...' if row['resource_type'] and len(row['resource_type']) > 15 else (row['resource_type'] or ''),
                        duration_str,
                        row['timestamp'].strftime('%m-%d %H:%M') if row['timestamp'] else ''
                    ])
                
                table = Table(table_data, colWidths=[1.2*inch, 1.5*inch, 1.2*inch, 1*inch, 1.6*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1976d2')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
            
            doc.build(story)
            pdf_buf.seek(0)
            return StreamingResponse(
                iter([pdf_buf.getvalue()]),
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename=user_activities_{timestamp}.pdf"},
            )
            
    except Exception as e:
        logger.exception(f"User activities export error: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {e!s}")


@router.get("/compliance-report")
@export_rate_limit
def export_compliance_report(
    request: Request,
    format: str = Query("csv", pattern="^(csv|excel|json|pdf|docx)$"),
    days: int = Query(30, ge=1, le=365, description="Days of data to include"),
    min_score: float = Query(0.0, ge=0.0, le=1.0, description="Minimum compliance score"),
    db: Session = Depends(get_session),
    current_user=Depends(require_role(["admin", "reviewer"])),
):
    """Export comprehensive compliance report"""
    try:
        from datetime import datetime, timedelta
        
        cutoff_date = datetime.now() - timedelta(days=days)
        
        # Get comprehensive compliance data
        query = (
            select(AIResult, FileUpload, Checklist, User)
            .join(FileUpload, AIResult.file_upload_id == FileUpload.id)
            .join(Checklist, AIResult.checklist_id == Checklist.id)
            .join(User, AIResult.user_id == User.id)
            .where(AIResult.created_at >= cutoff_date)
            .where(AIResult.score >= min_score)
        )
        
        results = db.exec(query).all()
        
        # Prepare compliance data
        data = []
        for ai_result, file_upload, checklist, user in results:
            # Determine compliance status
            compliance_status = "Compliant" if ai_result.score >= 0.7 else "Non-Compliant"
            risk_level = (
                "Low" if ai_result.score >= 0.8 
                else "Medium" if ai_result.score >= 0.6 
                else "High"
            )
            
            data.append({
                "assessment_id": ai_result.id,
                "checklist_id": checklist.id,
                "checklist_title": checklist.title,
                "file_id": file_upload.id,
                "filename": file_upload.filename,
                "user_id": user.id,
                "username": user.username,
                "user_role": user.role,
                "compliance_score": ai_result.score,
                "compliance_status": compliance_status,
                "risk_level": risk_level,
                "ai_model_version": ai_result.ai_model_version,
                "processing_time_ms": ai_result.processing_time_ms,
                "assessment_date": ai_result.created_at,
                "file_upload_date": file_upload.uploaded_at,
                "feedback_summary": (
                    ai_result.feedback[:200] + "..." if len(ai_result.feedback) > 200 
                    else ai_result.feedback
                ),
            })
        
        df = pd.DataFrame(data)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if format == "csv":
            csv_buf = StringIO()
            df.to_csv(csv_buf, index=False)
            csv_buf.seek(0)
            return StreamingResponse(
                iter([csv_buf.getvalue()]),
                media_type="text/csv",
                headers={"Content-Disposition": f"attachment; filename=compliance_report_{timestamp}.csv"},
            )
        elif format == "excel":
            excel_buf = BytesIO()
            df.to_excel(excel_buf, index=False, engine="openpyxl")
            excel_buf.seek(0)
            return StreamingResponse(
                iter([excel_buf.getvalue()]),
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": f"attachment; filename=compliance_report_{timestamp}.xlsx"},
            )
        elif format == "json":
            json_buf = StringIO()
            df.to_json(json_buf, orient="records", indent=2)
            json_buf.seek(0)
            return StreamingResponse(
                iter([json_buf.getvalue()]),
                media_type="application/json",
                headers={"Content-Disposition": f"attachment; filename=compliance_report_{timestamp}.json"},
            )
        elif format == "pdf":
            pdf_buf = BytesIO()
            doc = SimpleDocTemplate(pdf_buf, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.HexColor('#1976d2')
            )
            story.append(Paragraph("Compliance Report Export", title_style))
            story.append(Spacer(1, 12))
            
            # Summary
            summary_style = styles['Normal']
            story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", summary_style))
            story.append(Paragraph(f"Total Assessments: {len(data)}", summary_style))
            story.append(Paragraph(f"Data Period: Last {days} days", summary_style))
            story.append(Paragraph(f"Minimum Score Filter: {min_score}", summary_style))
            
            # Compliance statistics
            if data:
                compliant_count = len([d for d in data if d['compliance_status'] == 'Compliant'])
                compliance_rate = (compliant_count / len(data)) * 100 if data else 0
                story.append(Paragraph(f"Compliance Rate: {compliance_rate:.1f}% ({compliant_count}/{len(data)})", summary_style))
            
            story.append(Spacer(1, 20))
            
            # Data table
            if data:
                table_data = [['File', 'User', 'Score', 'Status', 'Risk', 'Date']]
                for row in data:
                    table_data.append([
                        row['filename'][:20] + '...' if len(row['filename']) > 20 else row['filename'],
                        row['username'][:15] + '...' if len(row['username']) > 15 else row['username'],
                        f"{row['compliance_score']:.1%}" if row['compliance_score'] is not None else 'N/A',
                        row['compliance_status'],
                        row['risk_level'],
                        row['assessment_date'].strftime('%Y-%m-%d') if row['assessment_date'] else ''
                    ])
                
                table = Table(table_data, colWidths=[1.8*inch, 1.2*inch, 0.8*inch, 1*inch, 0.8*inch, 1*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1976d2')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
            
            doc.build(story)
            pdf_buf.seek(0)
            return StreamingResponse(
                iter([pdf_buf.getvalue()]),
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename=compliance_report_{timestamp}.pdf"},
            )
            
    except Exception as e:
        logger.exception(f"Compliance report export error: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {e!s}")
