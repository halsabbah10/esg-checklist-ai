import csv
import logging
import os
from datetime import datetime, timezone
from io import BytesIO, StringIO
from typing import Optional

import aiofiles  # type: ignore[import-untyped]
import openpyxl  # type: ignore[import-untyped]
import pandas as pd
import pdfplumber
from docx import Document
from fastapi import APIRouter, Depends, File, HTTPException, Query, UploadFile, status
from fastapi.responses import StreamingResponse
from fpdf import FPDF  # type: ignore[import-untyped]
from sqlmodel import Session, select

from app.services.realtime_analytics import (
    realtime_analytics,
    track_ai_processing,
    track_file_upload,
)
from app.utils.ai import ai_score_text_with_gemini
from app.utils.email import send_ai_score_notification
from app.utils.file_security import generate_secure_filepath, validate_upload_file
from app.utils.notifications import notify_user

from ..auth import require_role
from ..config import get_settings
from ..database import get_session
from ..models import AIResult, Checklist, ChecklistItem, FileUpload
from ..schemas import ChecklistCreate, ChecklistItemRead, ChecklistRead

logger = logging.getLogger(__name__)

# Get centralized settings
settings = get_settings()

router = APIRouter(prefix="/checklists", tags=["checklists"])


@router.post("/", response_model=ChecklistRead)
def create_checklist(
    checklist: ChecklistCreate,
    db: Session = Depends(get_session),
    current_user=Depends(require_role("admin")),
):
    new_checklist = Checklist(
        title=checklist.title,
        description=checklist.description,
        created_by=current_user.id,
    )
    db.add(new_checklist)
    db.commit()
    db.refresh(new_checklist)

    # Ensure the checklist has an ID after database insertion
    if new_checklist.id is None:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to create checklist",
        )

    # Add items
    for item in checklist.items:
        new_item = ChecklistItem(
            checklist_id=new_checklist.id,  # Now guaranteed to be int
            question_text=item.question_text,
            weight=item.weight,
            category=item.category,
        )
        db.add(new_item)
    db.commit()
    return new_checklist


@router.get("/", response_model=list[ChecklistRead])
def list_checklists(
    db: Session = Depends(get_session),
    _current_user=Depends(require_role("auditor")),  # All roles can view
):
    return db.exec(select(Checklist)).all()


@router.get("/search")
def search_checklists(
    q: str = Query(..., description="Search query for checklist title or description"),
    category: Optional[str] = Query(None, description="Filter by category"),
    limit: int = Query(20, ge=1, le=100, description="Maximum number of results"),
    db: Session = Depends(get_session),
    _current_user=Depends(require_role("auditor")),
):
    """
    Search checklists by title, description, or category.
    Returns matching checklists with relevance scoring.
    """
    try:
        # Start with all checklists
        query = select(Checklist).where(Checklist.is_active is True)

        # Execute query to get all checklists
        checklists = db.exec(query).all()

        # Apply search filter in Python
        search_lower = q.lower()
        matching_checklists = []

        for checklist in checklists:
            relevance_score = 0

            # Title matching (higher weight)
            if search_lower in checklist.title.lower():
                relevance_score += 10

            # Description matching (medium weight)
            if checklist.description and search_lower in checklist.description.lower():
                relevance_score += 5

            # Category filter
            if category and hasattr(checklist, "category") and checklist.category != category:
                continue

            if relevance_score > 0:
                matching_checklists.append(
                    {
                        "id": checklist.id,
                        "title": checklist.title,
                        "description": checklist.description,
                        "created_by": checklist.created_by,
                        "created_at": checklist.created_at,
                        "is_active": checklist.is_active,
                        "relevance_score": relevance_score,
                    }
                )

        # Sort by relevance score (descending)
        matching_checklists.sort(
            key=lambda x: (
                x["relevance_score"] if isinstance(x["relevance_score"], (int, float)) else 0
            ),
            reverse=True,
        )

        # Apply limit
        matching_checklists = matching_checklists[:limit]

        return {
            "results": matching_checklists,
            "total_count": len(matching_checklists),
            "search_query": q,
            "category_filter": category,
        }

    except Exception as e:
        logger.exception(f"Search error: {e}")
        raise HTTPException(status_code=500, detail=f"Search failed: {e!s}")


@router.get("/{checklist_id}", response_model=ChecklistRead)
def get_checklist(
    checklist_id: int,
    db: Session = Depends(get_session),
    _current_user=Depends(require_role("auditor")),
):
    checklist = db.get(Checklist, checklist_id)
    if not checklist:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Checklist with ID {checklist_id} not found",
        )
    return checklist


@router.get("/{checklist_id}/items", response_model=list[ChecklistItemRead])
def get_checklist_items(
    checklist_id: int,
    db: Session = Depends(get_session),
    _current_user=Depends(require_role("auditor")),
):
    return db.exec(select(ChecklistItem).where(ChecklistItem.checklist_id == checklist_id)).all()


# Ensure upload directory exists (will be created by file_security module)
os.makedirs(settings.upload_path, exist_ok=True)


@router.post("/{checklist_id}/upload")
async def upload_file(
    checklist_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_session),
    current_user=Depends(require_role("auditor")),
):
    """
    Secure file upload with comprehensive validation
    """
    logger.info(f"User {current_user.id} uploading file for checklist {checklist_id}")

    try:
        # Validate checklist exists first
        checklist = db.exec(select(Checklist).where(Checklist.id == checklist_id)).first()
        if not checklist:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Checklist with ID {checklist_id} not found",
            )

        # Comprehensive file security validation
        secure_filename, file_extension = await validate_upload_file(file)

        # Generate secure file path
        secure_filepath = generate_secure_filepath(secure_filename, current_user.id, checklist_id)

        # Ensure directory exists
        secure_filepath.parent.mkdir(parents=True, exist_ok=True)

        # Save file securely using async I/O
        async with aiofiles.open(secure_filepath, "wb") as f:
            content = await file.read()
            await f.write(content)

        logger.info(f"File saved securely to: {secure_filepath}")

        # Store record in DB
        file_record = FileUpload(
            checklist_id=checklist_id,
            user_id=current_user.id,
            filename=secure_filename,
            filepath=str(secure_filepath),
        )
        db.add(file_record)
        db.commit()
        db.refresh(file_record)

        # Ensure the file record has an ID after database insertion
        if file_record.id is None:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to create file record",
            )

        # Extract text based on file extension using secure file path
        raw_text = ""
        try:
            if file_extension == "pdf":
                with pdfplumber.open(secure_filepath) as pdf:
                    raw_text = "\n".join([page.extract_text() or "" for page in pdf.pages])
            elif file_extension == "docx":
                doc = Document(str(secure_filepath))
                raw_text = "\n".join([p.text for p in doc.paragraphs])
            elif file_extension == "xlsx":
                wb = openpyxl.load_workbook(secure_filepath)
                text = []
                for ws in wb.worksheets:
                    for row in ws.iter_rows(values_only=True):
                        text.append(" ".join([str(cell) if cell else "" for cell in row]))
                raw_text = "\n".join(text)
            elif file_extension == "csv":
                with open(secure_filepath, "r", encoding="utf-8") as f:
                    reader = csv.reader(f)
                    raw_text = "\n".join([", ".join(row) for row in reader])
            elif file_extension == "txt":
                with open(secure_filepath, "r", encoding="utf-8", errors="ignore") as f:
                    raw_text = f.read()
            else:
                # Should not reach here due to validation, but fallback
                raise ValueError(f"Unsupported file extension: {file_extension}")
        except Exception as e:
            logger.exception(f"Error extracting text from {secure_filepath}: {e}")
            raw_text = f"Error extracting text: {e}"

        # AI/NLP scoring using Gemini with real-time tracking
        logger.info(f"Starting AI scoring for file: {secure_filename}")
        ai_start_time = datetime.now(timezone.utc)

        try:
            score, feedback = ai_score_text_with_gemini(raw_text)
            ai_end_time = datetime.now(timezone.utc)
            processing_time_ms = int((ai_end_time - ai_start_time).total_seconds() * 1000)

            logger.info(
                f"AI scoring completed - Score: {score:.3f}, Feedback length: {len(feedback)} chars"
            )

            # Track AI processing metrics
            track_ai_processing(
                db=db,
                user_id=current_user.id,
                session_id=f"upload_{file_record.id}",
                file_id=file_record.id,
                ai_score=score,
                processing_time_ms=processing_time_ms,
            )

        except Exception as e:
            logger.exception(f"AI scoring failed for file {secure_filename}: {e}")
            # Provide fallback score and feedback
            score = 0.5
            feedback = f"AI scoring temporarily unavailable. Error: {str(e)[:200]}..."

        # Truncate text if too long for database (TEXT can hold ~65k chars)
        max_text_length = 65000
        if len(raw_text) > max_text_length:
            raw_text = raw_text[:max_text_length] + "...[truncated]"

        if len(feedback) > max_text_length:
            feedback = feedback[:max_text_length] + "...[truncated]"

        # Store AI result in DB
        ai_result = AIResult(
            file_upload_id=file_record.id,  # Now guaranteed to be int
            checklist_id=checklist_id,
            user_id=current_user.id,
            raw_text=raw_text,
            score=score,
            feedback=feedback,
        )
        db.add(ai_result)
        db.commit()
        db.refresh(ai_result)

        # Send email notification asynchronously (best effort)
        try:
            send_ai_score_notification(
                user_email=current_user.email,
                filename=secure_filename,
                score=score,
                feedback=feedback,
                checklist_title=checklist.title,
            )
        except Exception as e:
            # Log error but don't fail the upload
            logger.exception(f"Failed to send email notification: {e}")

        # Send in-app notification for successful upload
        try:
            notify_user(
                db=db,
                user_id=current_user.id,
                title="File Upload Successful ✅",
                message=(
                    f"Your file '{secure_filename}' has been uploaded and analyzed. "
                    f"AI Score: {score:.3f}/1.0 ({score * 100:.1f}%)"
                ),
                link=f"/uploads/{file_record.id}",
                notification_type="success",
            )
        except Exception as e:
            # Log error but don't fail the upload
            logger.exception(f"Failed to send upload notification: {e}")

        # Track file upload completion with real-time analytics
        try:
            upload_end_time = datetime.now(timezone.utc)
            total_processing_time = int((upload_end_time - ai_start_time).total_seconds() * 1000)

            track_file_upload(
                db=db,
                user_id=current_user.id,
                session_id=f"upload_{file_record.id}",
                file_id=file_record.id,
                filename=secure_filename,
                processing_time_ms=total_processing_time,
            )

            # Track compliance metrics for analytics
            realtime_analytics.track_compliance_update(
                db=db,
                checklist_id=checklist_id,
                file_upload_id=file_record.id,
                compliance_score=score,
                risk_level="High" if score < 0.5 else "Medium" if score < 0.7 else "Low",
                recommendations=["Improve ESG documentation", "Enhance reporting quality"],
            )

        except Exception as e:
            # Log error but don't fail the upload
            logger.exception(f"Failed to track analytics: {e}")
            logger.exception(f"Failed to send upload notification: {e}")

        return {
            "detail": "File uploaded and AI scored",
            "file_id": file_record.id,
            "filename": secure_filename,
            "ai_score": score,
            "ai_feedback": feedback,
            "email_sent": True,  # Could be enhanced to track actual status
        }

    except HTTPException:
        # Re-raise HTTP exceptions (they're already properly handled)
        raise
    except Exception as e:
        logger.error(f"Unexpected error during file upload: {e}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="An unexpected error occurred during file upload",
        )


@router.get("/{checklist_id}/export", tags=["checklists"])
def export_checklist_results(
    checklist_id: int,
    export_format: str = "csv",
    db: Session = Depends(get_session),
    _current_user=Depends(require_role("admin")),
):
    # Query all uploads & AI results for this checklist
    results = db.exec(
        select(FileUpload, AIResult)
        .where(FileUpload.checklist_id == checklist_id)
        .where(FileUpload.id == AIResult.file_upload_id)
    ).all()

    # Prepare data
    data = []
    for upload, ai in results:
        data.append(
            {
                "file_id": upload.id,
                "filename": upload.filename,
                "user_id": upload.user_id,
                "ai_score": ai.score,
                "ai_feedback": ai.feedback,
                "uploaded_at": upload.uploaded_at,
            }
        )

    results_data = pd.DataFrame(data)

    # Ensure exports folder exists
    os.makedirs("exports", exist_ok=True)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    export_filename = f"exports/checklist_{checklist_id}_results_{timestamp}.{export_format}"

    # Export as CSV
    if export_format == "csv":
        results_data.to_csv(export_filename, index=False)
        buf = StringIO()
        results_data.to_csv(buf, index=False)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="text/csv",
            headers={
                "Content-Disposition": f"attachment; filename=checklist_{checklist_id}_results.csv"
            },
        )

    # Export as Excel
    if export_format == "excel":
        results_data.to_excel(export_filename, index=False)
        excel_buf = BytesIO()
        results_data.to_excel(excel_buf, index=False)
        excel_buf.seek(0)
        return StreamingResponse(
            excel_buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={
                "Content-Disposition": f"attachment; filename=checklist_{checklist_id}_results.xlsx"
            },
        )

    # Export as Word
    if export_format == "word":
        doc = Document()
        doc.add_heading(f"Checklist {checklist_id} Results", 0)
        table = doc.add_table(rows=1, cols=len(results_data.columns))
        hdr_cells = table.rows[0].cells
        for idx, column in enumerate(results_data.columns):
            hdr_cells[idx].text = column
        for row in results_data.itertuples(index=False):
            row_cells = table.add_row().cells
            for idx, value in enumerate(row):
                row_cells[idx].text = str(value)
        doc.save(export_filename)
        word_buf = BytesIO()
        doc.save(word_buf)
        word_buf.seek(0)
        return StreamingResponse(
            word_buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=checklist_{checklist_id}_results.docx"
            },
        )

    # Export as PDF
    if export_format == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=10)
        pdf.cell(0, 10, f"Checklist {checklist_id} Results", ln=True)
        # Table header
        for col in results_data.columns:
            pdf.cell(40, 10, col, border=1)
        pdf.ln()
        # Table rows
        for row in results_data.itertuples(index=False):
            for value in row:
                # Only print first 30 chars to keep things neat
                cell = str(value)[:30] if value is not None else ""
                pdf.cell(40, 10, cell, border=1)
            pdf.ln()
        pdf.output(export_filename)
        # Serve file
        with open(export_filename, "rb") as f:
            pdf_bytes = f.read()
        pdf_buf = BytesIO(pdf_bytes)
        pdf_buf.seek(0)
        return StreamingResponse(
            pdf_buf,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename=checklist_{checklist_id}_results.pdf"
            },
        )

    raise HTTPException(status_code=400, detail="Unsupported format")
